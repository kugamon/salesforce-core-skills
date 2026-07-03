#!/usr/bin/env python3
"""
Audit Report Generator

Takes scored JSON input files produced by the AI agent and generates
deterministic HTML, DOCX, XLSX, and JSON audit reports.

Usage:
    python generate_reports.py --input-dir audit_output --output-dir audit_output \
        --org-name "Acme Corp" --org-id 00D000000000001 --instance CS42

Dependencies:
    pip install openpyxl python-docx
"""

import argparse
import html
import json
import sys
from datetime import date
from pathlib import Path

# Optional deps — fail gracefully with clear message
try:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import docx
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ── Brand tokens (from references/report-template.md) ─────────────────────

BRAND_BLUE = "#417AE4"
BRAND_CYAN = "#14DDDD"
BODY_BG = "#F4F6F9"
BORDER = "#E0E4EA"
MUTED = "#6B7280"

SCORE_THRESHOLDS = [
    (80, "Excellent", "#E8FBF9", "#14DDDD"),
    (70, "Good", "#E9F7EF", "#27AE60"),
    (60, "Acceptable", "#FEF3CD", "#F39C12"),
    (40, "Needs Improvement", "#FEF9E7", "#E67E22"),
    (0, "Critical", "#FDE8E8", "#E74C3C"),
]

SEVERITY_COLORS = {
    "CRITICAL": ("#FDE8E8", "#E74C3C"),
    "HIGH": ("#FEF3CD", "#E67E22"),
    "MEDIUM": ("#EBF1FB", "#417AE4"),
    "LOW": ("#E9F7EF", "#27AE60"),
}


def score_rating(pct):
    """Return (rating_label, badge_bg, badge_text) for a percentage score."""
    for threshold, label, bg, text in SCORE_THRESHOLDS:
        if pct >= threshold:
            return label, bg, text
    return SCORE_THRESHOLDS[-1][1:]


def _esc(value):
    """HTML-escape a value, converting None to empty string."""
    return html.escape(str(value)) if value is not None else ""


# ── Input loading ───────────────────────────────────────────────────────────


INPUT_FILES = {
    "counts": "counts.json",
    "apex_scores": "apex_scores.json",
    "trigger_findings": "trigger_findings.json",
    "flow_scores": "flow_scores.json",
    "process_builders": "process_builders.json",
    "lwc_scores": "lwc_scores.json",
    "permission_findings": "permission_findings.json",
    "metadata_scores": "metadata_scores.json",
    "unused_fields": "unused_fields.json",
    "unused_objects": "unused_objects.json",
    "validation_rules": "validation_rules.json",
    "formula_fields": "formula_fields.json",
    "workflow_rules": "workflow_rules.json",
    "other_rules_findings": "other_rules_findings.json",
    "reports_dashboards": "reports_dashboards.json",
    "integrations": "integrations.json",
    "test_coverage": "test_coverage.json",
    "team_evaluation": "team_evaluation.json",
    "change_history": "change_history.json",
    "licensing": "licensing.json",
    "data_quality": "data_quality.json",
}

# Keys whose default empty value is a dict (not a list)
_DICT_DEFAULTS = {"counts", "test_coverage", "licensing", "team_evaluation",
                  "change_history", "reports_dashboards", "data_quality"}


def load_inputs(input_dir):
    """Load all JSON input files. Missing files become empty dicts/lists."""
    data = {}
    input_path = Path(input_dir)
    for key, filename in INPUT_FILES.items():
        fpath = input_path / filename
        if fpath.exists():
            with open(fpath, encoding="utf-8") as f:
                data[key] = json.load(f)
        else:
            data[key] = {} if key in _DICT_DEFAULTS else []
    return data


# ── Score computation ───────────────────────────────────────────────────────


def compute_domain_score(items, max_score_key="max_score", score_key="score", default_max=100):
    """Compute average percentage for a list of scored items."""
    if not items:
        return 0.0
    total_pct = 0.0
    for item in items:
        max_s = item.get(max_score_key, default_max)
        s = item.get(score_key, 0)
        total_pct += (s / max_s * 100) if max_s > 0 else 0
    return round(total_pct / len(items), 1)


def compute_summary(data):
    """Compute the full audit summary from loaded data."""
    counts = data["counts"]

    # Domain-specific max scores (Apex=150, Flows=110, LWC=165, Metadata=120)
    domain_config = [
        ("apex", "apex_scores", 150),
        ("flows", "flow_scores", 110),
        ("lwc", "lwc_scores", 165),
        ("metadata", "metadata_scores", 120),
    ]

    domain_scores = {}
    for domain, items_key, default_max in domain_config:
        items = data.get(items_key, [])
        domain_scores[domain] = compute_domain_score(
            items, max_score_key="max_score", score_key="score", default_max=default_max
        )

    # Overall = average of domain percentages (only domains with data)
    active_scores = []
    for domain, items_key, _default_max in domain_config:
        if data.get(items_key):
            active_scores.append(domain_scores[domain])

    overall = round(sum(active_scores) / len(active_scores), 1) if active_scores else 0.0

    # Below-threshold counts (<70%)
    below_threshold = {}
    for domain, items_key, default_max in domain_config:
        items = data.get(items_key, [])
        count = 0
        for item in items:
            max_s = item.get("max_score", default_max)
            s = item.get("score", 0)
            if max_s > 0 and (s / max_s * 100) < 70:
                count += 1
        below_threshold[domain] = count

    # Severity rollup from permission findings and all declarative-logic findings
    severity_counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
    for finding in data.get("permission_findings", []):
        sev = finding.get("severity", "LOW").upper()
        if sev in severity_counts:
            severity_counts[sev] += 1
    for source_key in ("validation_rules", "formula_fields", "workflow_rules", "other_rules_findings"):
        for item in data.get(source_key, []):
            for finding in item.get("findings", []):
                sev = finding.get("severity", "LOW").upper()
                if sev in severity_counts:
                    severity_counts[sev] += 1
    for source_key in ("unused_fields", "unused_objects"):
        for item in data.get(source_key, []):
            sev = item.get("severity", "LOW").upper()
            if sev in severity_counts:
                severity_counts[sev] += 1

    # Top issues by domain
    top_issues = {}
    for domain, items_key, _default_max in domain_config:
        issue_counts = {}
        for item in data.get(items_key, []):
            for issue in item.get("issues", []):
                label = issue if isinstance(issue, str) else issue.get("message", str(issue))
                issue_counts[label] = issue_counts.get(label, 0) + 1
        sorted_issues = sorted(issue_counts.items(), key=lambda x: -x[1])
        top_issues[domain] = sorted_issues[:3]

    # Test coverage summary
    tc = data.get("test_coverage", {})
    test_coverage_summary = {
        "org_wide_pct": tc.get("org_wide_pct") or 0.0,
        "classes_with_zero_coverage": sum(
            1 for c in tc.get("classes", []) if c.get("coverage_pct", 0) == 0
        ),
        "total_classes_tracked": len(tc.get("classes", [])),
    }

    # Licensing summary
    lic = data.get("licensing", {})
    licensing_summary = {
        "user_licenses": len(lic.get("user_licenses", [])),
        "psl": len(lic.get("permission_set_licenses", [])),
        "package_licenses": len(lic.get("package_licenses", [])),
        "underutilised": sum(
            1 for lic_item in lic.get("user_licenses", [])
            + lic.get("permission_set_licenses", [])
            if lic_item.get("utilization_pct", 100) < 10
        ) + sum(
            1 for pkg in lic.get("package_licenses", [])
            if pkg.get("allowed", 0) > 0
            and (pkg.get("used", 0) / pkg["allowed"] * 100) < 10
        ),
    }

    # Team summary
    team = data.get("team_evaluation", {})
    team_summary = {
        "active_users": team.get("active_users", 0),
        "inactive_90d": sum(
            1 for u in team.get("users", [])
            if (u.get("days_since_login") or 0) > 90
        ),
    }

    # Reports & Dashboards summary
    rd = data.get("reports_dashboards", {})
    reports_dashboards_summary = {
        "reports": len(rd.get("reports", [])),
        "dashboards": len(rd.get("dashboards", [])),
        "stale": sum(1 for r in rd.get("reports", []) if r.get("is_stale"))
        + sum(1 for d in rd.get("dashboards", []) if d.get("is_stale")),
    }

    # Change history summary
    ch = data.get("change_history", {})
    change_history_summary = {
        "deployment_success_rate": ch.get("deployment_success_rate", 0),
        "total_deployments": len(ch.get("deployments", [])),
        "audit_trail_entries": len(ch.get("audit_trail", [])),
    }

    # Collect severity counts from new domains too
    for source_key in ("integrations",):
        for item in data.get(source_key, []):
            for finding in item.get("findings", []):
                sev = finding.get("severity", "LOW").upper()
                if sev in severity_counts:
                    severity_counts[sev] += 1
    for source_key in ("test_coverage", "licensing", "team_evaluation",
                       "change_history", "reports_dashboards", "data_quality"):
        obj = data.get(source_key, {})
        for finding in obj.get("findings", []):
            sev = finding.get("severity", "LOW").upper()
            if sev in severity_counts:
                severity_counts[sev] += 1

    return {
        "org_name": counts.get("org_name", ""),
        "org_id": counts.get("org_id", ""),
        "instance": counts.get("instance", ""),
        "overall_score": overall,
        "overall_rating": score_rating(overall)[0],
        "domain_scores": domain_scores,
        "counts": counts,
        "below_threshold": below_threshold,
        "severity_counts": severity_counts,
        "top_issues": {k: [{"issue": i, "count": c} for i, c in v] for k, v in top_issues.items()},
        "test_coverage": test_coverage_summary,
        "licensing": licensing_summary,
        "team": team_summary,
        "reports_dashboards": reports_dashboards_summary,
        "change_history": change_history_summary,
        "data_quality": {
            "objects_checked": len(data.get("data_quality", {}).get("objects", [])),
            "findings_count": len(data.get("data_quality", {}).get("findings", [])),
        },
    }


# ── Hardcoded values cross-cutting helper ──────────────────────────────────


def _collect_hardcoded_values(data):
    """Gather all hardcoded-value findings across every declarative source.

    Returns a list of dicts with keys: component_type, name, severity, message.
    """
    rows = []
    source_map = [
        ("validation_rules", "Validation Rule"),
        ("formula_fields", "Formula Field"),
        ("workflow_rules", "Workflow Rule"),
        ("other_rules_findings", None),  # type comes from the item itself
    ]
    for source_key, default_type in source_map:
        for item in data.get(source_key, []):
            comp_type = default_type or item.get("type", "Unknown")
            for finding in item.get("findings", []):
                msg = finding.get("message", "")
                if "hardcoded" in msg.lower():
                    rows.append({
                        "component_type": comp_type,
                        "name": item.get("name", ""),
                        "severity": finding.get("severity", "MEDIUM"),
                        "message": msg,
                    })
    return rows


# ── HTML report ─────────────────────────────────────────────────────────────


def _severity_badge_html(severity):
    bg, fg = SEVERITY_COLORS.get(severity.upper(), ("#EBF1FB", "#417AE4"))
    return (
        f'<span style="background:{bg};color:{fg};padding:2px 8px;'
        f'border-radius:10px;font-size:11px;font-weight:600">{_esc(severity)}</span>'
    )


def _score_badge_html(score, max_score):
    pct = (score / max_score * 100) if max_score > 0 else 0
    _, bg, fg = score_rating(pct)
    return (
        f'<span style="background:{bg};color:{fg};padding:2px 8px;'
        f'border-radius:10px;font-size:11px;font-weight:600">'
        f"{score}/{max_score}</span>"
    )


def _table_html(headers, rows):
    """Build an HTML table with branded headers."""
    parts = ["<table>", "<thead><tr>"]
    for h in headers:
        parts.append(f"<th>{_esc(h)}</th>")
    parts.append("</tr></thead><tbody>")
    for row in rows:
        parts.append("<tr>")
        for cell in row:
            parts.append(f"<td>{cell}</td>")  # cell may contain HTML badges
        parts.append("</tr>")
    parts.append("</tbody></table>")
    return "\n".join(parts)


def _findings_html(findings):
    """Render a list of findings with severity left-border strips."""
    if not findings:
        return "<p>No findings.</p>"
    parts = []
    for f in findings:
        sev = f.get("severity", "MEDIUM").lower()
        css_class = {
            "critical": "critical",
            "high": "warning",
            "medium": "info",
            "low": "positive",
        }.get(sev, "info")
        parts.append(
            f'<div class="finding {css_class}">'
            f'<span class="finding-badge">{_esc(f.get("severity", "MEDIUM"))}</span> '
            f'{_esc(f.get("message", f.get("finding", "")))}'
            f"</div>"
        )
    return "\n".join(parts)


def _recommendations_list(data):
    """Build the top-10 recommendations sorted by priority.

    Returns a list of plain-text recommendation strings (no HTML markup).
    CRITICAL permission findings rank highest so they aren't buried by
    low-scoring domain items.
    """
    scored_recs = []  # list of (priority, text)

    # Below-threshold domain items: priority based on how far below 70%
    for domain, items_key, default_max in [
        ("Apex", "apex_scores", 150),
        ("Flows", "flow_scores", 110),
        ("LWC", "lwc_scores", 165),
        ("Metadata", "metadata_scores", 120),
    ]:
        for item in data.get(items_key, []):
            s = item.get("score", 0)
            max_s = item.get("max_score", default_max)
            pct = (s / max_s * 100) if max_s > 0 else 0
            if pct < 70:
                name = item.get("name", "Unknown")
                top_issue = ""
                issues = item.get("issues", [])
                if issues:
                    top_issue = issues[0] if isinstance(issues[0], str) else issues[0].get("message", "")
                # Priority: 20-69 based on percentage (lower = higher rank)
                priority = 20 + pct
                issue_suffix = f" {top_issue}" if top_issue else ""
                scored_recs.append((
                    priority,
                    f"[{domain}] Fix {name} (score {s}/{max_s}).{issue_suffix}",
                ))

    # Permission findings: CRITICAL=0, HIGH=5, MEDIUM=30, LOW=50
    sev_priority = {"CRITICAL": 0, "HIGH": 5, "MEDIUM": 30, "LOW": 50}
    for f in data.get("permission_findings", []):
        sev = f.get("severity", "LOW").upper()
        scored_recs.append((
            sev_priority.get(sev, 50),
            f"[Permissions] {f.get('message', f.get('finding', ''))}",
        ))

    # Legacy automation: moderate priority
    pb_count = len(data.get("process_builders", []))
    wr_count = len(data.get("workflow_rules", []))
    if pb_count:
        scored_recs.append((35, f"[Automation] Migrate {pb_count} active Process Builder(s) to Flow"))
    if wr_count:
        scored_recs.append((35, f"[Automation] Migrate {wr_count} Workflow Rule(s) to Flow"))

    # Hardcoded-value findings from formulas, validation rules, workflow rules, other rules
    for source_key, label in [
        ("validation_rules", "Validation Rules"),
        ("formula_fields", "Formula Fields"),
        ("workflow_rules", "Workflow Rules"),
        ("other_rules_findings", "Declarative Logic"),
    ]:
        for item in data.get(source_key, []):
            for finding in item.get("findings", []):
                sev = finding.get("severity", "LOW").upper()
                msg = finding.get("message", "")
                if "hardcoded" in msg.lower():
                    scored_recs.append((
                        sev_priority.get(sev, 50),
                        f"[{label}] {item.get('name', 'Unknown')}: {msg}",
                    ))

    # Unused fields and objects
    unused_fields = [f for f in data.get("unused_fields", []) if f.get("category") == "Unused"]
    unused_objects = [o for o in data.get("unused_objects", []) if o.get("category") == "Unused"]
    if unused_fields:
        scored_recs.append((
            15,
            f"[Data Model] Remove {len(unused_fields)} unused custom field(s) (no data, no references)",
        ))
    if unused_objects:
        scored_recs.append((
            10,
            f"[Data Model] Remove {len(unused_objects)} unused custom object(s) (no records, no references)",
        ))

    scored_recs.sort(key=lambda x: x[0])
    return [text for _, text in scored_recs[:10]]


def _recommendations_html(data):
    """Build the top-10 recommendations section as HTML."""
    recs = _recommendations_list(data)
    if not recs:
        return "<p>No recommendations — org is in great shape.</p>"

    parts = []
    for i, rec in enumerate(recs, 1):
        parts.append(
            f'<div style="display:flex;gap:12px;align-items:flex-start;margin-bottom:10px">'
            f'<div class="rec-num">{i}</div>'
            f"<div>{_esc(rec)}</div></div>"
        )
    return "\n".join(parts)


def generate_html(data, summary, org_name, org_id, instance, run_date, output_path):
    """Generate the self-contained HTML audit report."""
    overall = summary["overall_score"]
    rating = summary["overall_rating"]
    counts = summary["counts"]

    # Core CSS is embedded below (from references/report-template.md)

    sections = []

    # ── Executive Summary ──
    inv_lines = []
    inv_map = [
        ("apex_classes", "Apex Classes"),
        ("apex_triggers", "Apex Triggers"),
        ("active_flows", "Active Flows"),
        ("process_builders", "Process Builders"),
        ("lwc_bundles", "LWC Components"),
        ("custom_objects", "Custom Objects"),
        ("validation_rules", "Validation Rules"),
        ("formula_fields", "Formula Fields"),
        ("workflow_rules", "Workflow Rules"),
        ("approval_processes", "Approval Processes"),
        ("permission_sets", "Permission Sets"),
        ("permission_set_groups", "Permission Set Groups"),
        ("profiles", "Profiles"),
    ]
    for key, label in inv_map:
        val = counts.get(key, 0)
        inv_lines.append(f"<li><strong>{_esc(val)}</strong> {label}</li>")
    # Unused fields/objects summary counts
    unused_f = data.get("unused_fields", [])
    unused_o = data.get("unused_objects", [])
    if unused_f:
        inv_lines.append(
            f"<li><strong>{len(unused_f)}</strong> Unused/Empty/Unreferenced Custom Fields</li>"
        )
    if unused_o:
        inv_lines.append(
            f"<li><strong>{len(unused_o)}</strong> Unused/Empty/Unreferenced Custom Objects</li>"
        )
    sections.append(
        f'<div class="card"><h2>Executive Summary</h2>'
        f'<ul>{"".join(inv_lines)}</ul></div>'
    )

    # ── Domain score cards ──
    domain_cards = []
    for domain, label, _max_s in [
        ("apex", "Apex", 150),
        ("flows", "Flows", 110),
        ("lwc", "LWC", 165),
        ("metadata", "Metadata", 120),
    ]:
        pct = summary["domain_scores"].get(domain, 0)
        r, bg, fg = score_rating(pct)
        below = summary["below_threshold"].get(domain, 0)
        domain_cards.append(
            f'<div class="score-card">'
            f'<div class="score-card-label">{label}</div>'
            f'<div style="font-size:28px;font-weight:700;color:{fg}">{pct}%</div>'
            f'<div style="font-size:11px;color:{MUTED}">{r}</div>'
            f'<div style="font-size:11px;color:{MUTED}">{below} below 70%</div>'
            f"</div>"
        )
    sections.append(
        f'<div class="card"><h2>Domain Scores</h2>'
        f'<div style="display:flex;gap:16px;flex-wrap:wrap">{"".join(domain_cards)}</div></div>'
    )

    # ── Apex Classes ──
    if data.get("apex_scores"):
        rows = []
        for item in sorted(data["apex_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 150)),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Apex Classes</h2>'
            f'{_table_html(["Name", "Score", "Top Issues"], rows)}</div>'
        )

    # ── Apex Triggers ──
    if data.get("trigger_findings"):
        rows = []
        for item in sorted(data["trigger_findings"], key=lambda x: x.get("name", "")):
            findings_str = "; ".join(
                f.get("message", f.get("finding", ""))
                for f in item.get("findings", [])[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in item.get("findings", [])),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            )
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(item.get("events", "")),
                _severity_badge_html(sev),
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Apex Triggers</h2>'
            f'{_table_html(["Name", "Object", "Events", "Severity", "Findings"], rows)}</div>'
        )

    # ── Flows ──
    if data.get("flow_scores"):
        rows = []
        for item in sorted(data["flow_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("process_type", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 110)),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Flows</h2>'
            f'{_table_html(["Name", "Process Type", "Score", "Top Issues"], rows)}</div>'
        )

    # ── Process Builders ──
    if data.get("process_builders"):
        rows = []
        for item in sorted(data["process_builders"], key=lambda x: x.get("name", "")):
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(str(item.get("criteria_count", ""))),
                _esc(item.get("actions_summary", "")),
                _severity_badge_html(item.get("migration_priority", "HIGH")),
            ])
        sections.append(
            f'<div class="card"><h2>Process Builders</h2>'
            f'{_table_html(["Name", "Object", "Criteria", "Actions", "Migration Priority"], rows)}</div>'
        )

    # ── LWC ──
    if data.get("lwc_scores"):
        rows = []
        for item in sorted(data["lwc_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 165)),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Lightning Web Components</h2>'
            f'{_table_html(["Name", "Score", "Top Issues"], rows)}</div>'
        )

    # ── Permissions ──
    if data.get("permission_findings"):
        sorted_findings = sorted(
            data["permission_findings"],
            key=lambda x: (
                {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}.get(
                    x.get("severity", "LOW").upper(), 4
                ),
                x.get("name", ""),
            ),
        )
        sections.append(
            f'<div class="card"><h2>Profiles &amp; Permissions</h2>'
            f"{_findings_html(sorted_findings)}</div>"
        )

    # ── Metadata / Data Model ──
    if data.get("metadata_scores"):
        rows = []
        for item in sorted(data["metadata_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 120)),
                _esc(str(item.get("field_count", ""))),
                _esc(str(item.get("relationship_count", ""))),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Data Model</h2>'
            f'{_table_html(["Object", "Score", "Fields", "Relationships", "Top Issues"], rows)}</div>'
        )

    # ── Unused Fields & Objects ──
    if data.get("unused_fields") or data.get("unused_objects"):
        uf_parts = []
        if data.get("unused_fields"):
            uf_rows = []
            for item in sorted(data["unused_fields"], key=lambda x: (
                {"HIGH": 0, "MEDIUM": 1, "LOW": 2}.get(x.get("severity", "LOW").upper(), 3),
                x.get("object", ""), x.get("field", ""),
            )):
                refs = ", ".join(item.get("referenced_in") or []) or "None"
                has_data = item.get("has_data")
                has_data_label = "Unknown" if has_data is None else ("Yes" if has_data else "No")
                uf_rows.append([
                    _esc(item.get("object", "")),
                    _esc(item.get("field", "")),
                    _esc(item.get("data_type", "")),
                    _esc(has_data_label),
                    _esc(refs),
                    _esc(item.get("category", "")),
                    _severity_badge_html(item.get("severity", "LOW")),
                ])
            uf_parts.append(
                f"<h3>Unused Fields ({len(data['unused_fields'])})</h3>"
                + _table_html(
                    ["Object", "Field", "Data Type", "Has Data", "Referenced In", "Category", "Severity"],
                    uf_rows,
                )
            )
        if data.get("unused_objects"):
            uo_rows = []
            for item in sorted(data["unused_objects"], key=lambda x: (
                {"HIGH": 0, "MEDIUM": 1, "LOW": 2}.get(x.get("severity", "LOW").upper(), 3),
                x.get("object", ""),
            )):
                refs = ", ".join(item.get("referenced_in") or []) or "None"
                uo_rows.append([
                    _esc(item.get("object", "")),
                    _esc(str(item.get("record_count", 0))),
                    _esc(refs),
                    _esc(item.get("category", "")),
                    _severity_badge_html(item.get("severity", "LOW")),
                ])
            uf_parts.append(
                f"<h3>Unused Objects ({len(data['unused_objects'])})</h3>"
                + _table_html(
                    ["Object", "Record Count", "Referenced In", "Category", "Severity"],
                    uo_rows,
                )
            )
        sections.append(
            '<div class="card"><h2>Unused Fields &amp; Objects</h2>'
            + "\n".join(uf_parts) + "</div>"
        )

    # ── Validation Rules ──
    if data.get("validation_rules"):
        rows = []
        for item in sorted(data["validation_rules"], key=lambda x: x.get("name", "")):
            findings_str = "; ".join(
                f.get("message", f.get("finding", ""))
                for f in item.get("findings", [])[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in item.get("findings", [])),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if item.get("findings") else "LOW"
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(str(item.get("active", ""))),
                _severity_badge_html(sev),
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Validation Rules</h2>'
            f'{_table_html(["Name", "Object", "Active", "Severity", "Findings"], rows)}</div>'
        )

    # ── Formula Fields ──
    if data.get("formula_fields"):
        rows = []
        for item in sorted(data["formula_fields"], key=lambda x: x.get("name", "")):
            findings_str = "; ".join(
                f.get("message", f.get("finding", ""))
                for f in item.get("findings", [])[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in item.get("findings", [])),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if item.get("findings") else "—"
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(item.get("data_type", "")),
                _esc(str(item.get("formula_length", ""))),
                _severity_badge_html(sev) if sev != "—" else "—",
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Formula Fields</h2>'
            f'{_table_html(["Name", "Object", "Data Type", "Length", "Severity", "Findings"], rows)}</div>'
        )

    # ── Workflow Rules ──
    if data.get("workflow_rules"):
        rows = []
        for item in sorted(data["workflow_rules"], key=lambda x: x.get("name", "")):
            findings_str = "; ".join(
                f.get("message", "")
                for f in item.get("findings", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(item.get("action_types", "")),
                _severity_badge_html(item.get("migration_priority", "HIGH")),
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Workflow Rules</h2>'
            f'{_table_html(["Name", "Object", "Action Types", "Priority", "Formula Findings"], rows)}</div>'
        )

    # ── Other Declarative Logic ──
    if data.get("other_rules_findings"):
        rows = []
        for item in sorted(data["other_rules_findings"], key=lambda x: x.get("name", "")):
            findings_str = "; ".join(
                f.get("message", "")
                for f in item.get("findings", [])[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in item.get("findings", [])),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if item.get("findings") else "—"
            rows.append([
                _esc(item.get("type", "")),
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _severity_badge_html(sev) if sev != "—" else "—",
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Other Declarative Logic</h2>'
            f'{_table_html(["Type", "Name", "Object", "Severity", "Findings"], rows)}</div>'
        )

    # ── Hardcoded Values Summary ──
    hv_rows = _collect_hardcoded_values(data)
    if hv_rows:
        table_rows = []
        for hv in hv_rows:
            table_rows.append([
                _esc(hv["component_type"]),
                _esc(hv["name"]),
                _severity_badge_html(hv["severity"]),
                _esc(hv["message"]),
            ])
        sections.append(
            f'<div class="card"><h2>Hardcoded Values Summary</h2>'
            f'<p>{len(hv_rows)} hardcoded value(s) found across all declarative logic.</p>'
            f'{_table_html(["Component Type", "Name", "Severity", "Finding"], table_rows)}</div>'
        )
    else:
        sections.append(
            '<div class="card"><h2>Hardcoded Values Summary</h2>'
            "<p>No hardcoded values detected in formulas or declarative logic.</p></div>"
        )

    # ── Reports & Dashboards ──
    rd = data.get("reports_dashboards", {})
    rd_reports = rd.get("reports", [])
    rd_dashboards = rd.get("dashboards", [])
    if rd_reports or rd_dashboards:
        rd_rows = []
        for item in rd_reports:
            rd_rows.append([
                _esc(item.get("name", "")),
                "Report",
                _esc(item.get("folder", "")),
                _esc(item.get("last_run_date") or "Never"),
                "Yes" if item.get("is_stale") else "No",
            ])
        for item in rd_dashboards:
            rd_rows.append([
                _esc(item.get("name", "")),
                "Dashboard",
                _esc(item.get("folder", "")),
                _esc(item.get("last_viewed_date") or "Never"),
                "Yes" if item.get("is_stale") else "No",
            ])
        sections.append(
            f'<div class="card"><h2>Reports &amp; Dashboards</h2>'
            f'{_table_html(["Name", "Type", "Folder", "Last Activity", "Stale"], rd_rows)}</div>'
        )

    # ── Integrations ──
    if data.get("integrations"):
        int_rows = []
        for item in sorted(data["integrations"], key=lambda x: x.get("type", "")):
            findings_str = "; ".join(
                f.get("message", "") for f in item.get("findings", [])[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in item.get("findings", [])),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if item.get("findings") else "—"
            int_rows.append([
                _esc(item.get("type", "")),
                _esc(item.get("name", "")),
                _esc(item.get("endpoint", "") or "—"),
                _severity_badge_html(sev) if sev != "—" else "—",
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Integration Analysis</h2>'
            f'{_table_html(["Type", "Name", "Endpoint", "Severity", "Findings"], int_rows)}</div>'
        )

    # ── Test Coverage ──
    tc = data.get("test_coverage", {})
    tc_classes = tc.get("classes", [])
    if tc_classes or tc.get("org_wide_pct") is not None:
        org_pct = tc.get("org_wide_pct") or 0.0
        _, tc_bg, tc_fg = score_rating(org_pct)
        tc_rows = []
        for item in sorted(tc_classes, key=lambda x: x.get("coverage_pct", 0)):
            tc_rows.append([
                _esc(item.get("name", "")),
                str(item.get("lines_covered", 0)),
                str(item.get("lines_uncovered", 0)),
                f'{item.get("coverage_pct", 0):.1f}%',
            ])
        sections.append(
            f'<div class="card"><h2>Test Coverage</h2>'
            f'<p>Org-wide coverage: <span style="background:{tc_bg};color:{tc_fg};'
            f'padding:2px 8px;border-radius:10px;font-weight:600">{org_pct:.1f}%</span></p>'
            f'{_table_html(["Class/Trigger", "Lines Covered", "Lines Uncovered", "Coverage %"], tc_rows)}</div>'
        )

    # ── Licensing ──
    lic = data.get("licensing", {})
    lic_all = lic.get("user_licenses", []) + lic.get("permission_set_licenses", []) + lic.get("package_licenses", [])
    if lic_all:
        lic_rows = []
        for item in lic.get("user_licenses", []):
            lic_rows.append([
                _esc(item.get("name", "")), "User License",
                str(item.get("total", "")), str(item.get("used", "")),
                str(item.get("available", "")), f'{item.get("utilization_pct", 0):.0f}%',
            ])
        for item in lic.get("permission_set_licenses", []):
            lic_rows.append([
                _esc(item.get("label", item.get("name", ""))), "PS License",
                str(item.get("total", "")), str(item.get("used", "")),
                str(item.get("available", "")), f'{item.get("utilization_pct", 0):.0f}%',
            ])
        for item in lic.get("package_licenses", []):
            lic_rows.append([
                _esc(item.get("namespace", "")), "Package License",
                str(item.get("allowed", "")), str(item.get("used", "")),
                str(item.get("allowed", 0) - item.get("used", 0)),
                f'{(item.get("used", 0) / item.get("allowed", 1) * 100) if item.get("allowed", 0) > 0 else 0:.0f}%',
            ])
        sections.append(
            f'<div class="card"><h2>Licensing Analysis</h2>'
            f'{_table_html(["Name", "Type", "Total", "Used", "Available", "Utilization"], lic_rows)}'
            f'{_findings_html(lic.get("findings", []))}</div>'
        )

    # ── Team Evaluation ──
    team = data.get("team_evaluation", {})
    team_users = team.get("users", [])
    if team_users:
        team_rows = []
        for item in sorted(team_users, key=lambda x: x.get("days_since_login") or 0, reverse=True)[:50]:
            team_rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("profile", "")),
                _esc(item.get("role", "") or "—"),
                _esc(item.get("last_login", "Never")),
                str(item.get("days_since_login", "—")),
            ])
        sections.append(
            f'<div class="card"><h2>Team Evaluation</h2>'
            f'<p>{team.get("active_users", 0)} active users</p>'
            f'{_table_html(["Name", "Profile", "Role", "Last Login", "Days Inactive"], team_rows)}'
            f'{_findings_html(team.get("findings", []))}</div>'
        )

    # ── Change History ──
    ch = data.get("change_history", {})
    ch_deployments = ch.get("deployments", [])
    if ch_deployments or ch.get("audit_trail"):
        ch_rows = []
        for item in ch_deployments[:20]:
            ch_rows.append([
                _esc(item.get("date", "")),
                _esc(item.get("user", "")),
                _esc(item.get("status", "")),
                str(item.get("components_deployed", "")),
                str(item.get("component_errors", "")),
            ])
        dep_rate = ch.get("deployment_success_rate", 0)
        sections.append(
            f'<div class="card"><h2>Change History</h2>'
            f'<p>Deployment success rate: <strong>{dep_rate:.0f}%</strong> '
            f'({len(ch_deployments)} deployments)</p>'
            f'{_table_html(["Date", "User", "Status", "Components", "Errors"], ch_rows)}'
            f'{_findings_html(ch.get("findings", []))}</div>'
        )

    # ── Recommendations ──
    sections.append(
        f'<div class="card"><h2>Recommendations</h2>'
        f"{_recommendations_html(data)}</div>"
    )

    body = "\n".join(sections)

    report = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Salesforce Org Audit Report — {_esc(org_name)}</title>
<style>
:root {{
  --brand-blue: {BRAND_BLUE};
  --brand-cyan: {BRAND_CYAN};
}}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
       background: {BODY_BG}; color: #1a1a2e; line-height: 1.5; }}

.banner {{
  background: {BRAND_BLUE};
  padding: 28px 40px 32px;
  display: flex; align-items: center; gap: 24px;
}}
.banner-text {{ flex: 1; }}
.banner-title {{ font-size: 26px; font-weight: 700; color: #fff; letter-spacing: -0.3px; }}
.banner-subtitle {{ font-size: 13px; color: rgba(255,255,255,.80); margin-top: 4px; }}
.banner-score {{ text-align: right; flex-shrink: 0; }}
.banner-score-value {{ font-size: 48px; font-weight: 800; color: #fff; line-height: 1; }}
.banner-score-label {{ font-size: 11px; text-transform: uppercase; letter-spacing: 1px;
                       color: rgba(255,255,255,.75); margin-top: 4px; }}
.banner-score-rating {{
  display: inline-block; margin-top: 6px;
  background: rgba(255,255,255,.20); border: 1px solid rgba(255,255,255,.35);
  border-radius: 20px; padding: 2px 12px;
  font-size: 12px; font-weight: 600; color: #fff;
}}

.card {{ background: #fff; border: 1px solid {BORDER}; border-radius: 8px;
         padding: 24px; margin: 16px 40px; }}
.card h2 {{ color: {BRAND_BLUE}; font-size: 18px; margin-bottom: 16px; }}

table {{ width: 100%; border-collapse: collapse; margin-top: 8px; }}
thead th {{
  background: {BRAND_BLUE}; color: #fff; font-size: 11px; font-weight: 700;
  text-transform: uppercase; letter-spacing: 0.6px;
  padding: 10px 14px; text-align: left;
}}
tbody td {{ padding: 10px 14px; border-bottom: 1px solid {BORDER}; font-size: 13px; }}
tbody tr:hover {{ background: #f8f9fb; }}

.score-card {{
  background: #fff; border: 1px solid {BORDER}; border-radius: 8px;
  padding: 16px 20px; min-width: 140px; text-align: center;
}}
.score-card-label {{ font-size: 12px; font-weight: 600; color: {MUTED};
                     text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 4px; }}

.rec-num {{
  width: 28px; height: 28px; border-radius: 50%; background: {BRAND_BLUE};
  color: #fff; font-size: 13px; font-weight: 700;
  display: flex; align-items: center; justify-content: center; flex-shrink: 0;
}}

.finding {{ border-left: 4px solid; border-radius: 6px; padding: 12px 16px; margin-bottom: 8px;
            background: #fff; }}
.finding.critical {{ border-left-color: #E74C3C; }}
.finding.warning {{ border-left-color: #E67E22; }}
.finding.info {{ border-left-color: {BRAND_BLUE}; }}
.finding.positive {{ border-left-color: #27AE60; }}
.finding-badge {{
  display: inline-block; padding: 2px 8px; border-radius: 10px;
  font-size: 11px; font-weight: 600; margin-right: 8px;
}}
.finding.critical .finding-badge {{ background: #FDE8E8; color: #E74C3C; }}
.finding.warning .finding-badge {{ background: #FEF3CD; color: #E67E22; }}
.finding.info .finding-badge {{ background: #EBF1FB; color: {BRAND_BLUE}; }}
.finding.positive .finding-badge {{ background: #E9F7EF; color: #27AE60; }}

.footer {{
  background: {BRAND_BLUE};
  padding: 20px 40px; text-align: center;
  font-size: 12px; color: rgba(255,255,255,.80); margin-top: 24px;
}}
.footer a {{ color: #fff; }}
</style>
</head>
<body>
<div class="banner">
  <div class="banner-text">
    <div class="banner-title">Salesforce Org Audit Report</div>
    <div class="banner-subtitle">
      {_esc(org_name)} &middot; Org ID: {_esc(org_id)} &middot;
      Instance: {_esc(instance)} &middot; {_esc(run_date)}
    </div>
  </div>
  <div class="banner-score">
    <div class="banner-score-value">{overall:.0f}</div>
    <div class="banner-score-label">out of 100</div>
    <div class="banner-score-rating">{_esc(rating)}</div>
  </div>
</div>
{body}
<div class="footer">
  Generated by Salesforce Core Skills Audit Engine
  &nbsp;&middot;&nbsp; {_esc(run_date)} &nbsp;&middot;&nbsp;
  Org: {_esc(org_name)} ({_esc(org_id)})
</div>
</body>
</html>"""

    Path(output_path).write_text(report, encoding="utf-8")
    return output_path


# ── DOCX report ─────────────────────────────────────────────────────────────


def _hex_to_rgb(hex_color):
    """Convert '#RRGGBB' or 'RRGGBB' to RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def generate_docx(data, summary, org_name, org_id, instance, run_date, output_path):
    """Generate the Word audit report."""
    if not HAS_DOCX:
        print("WARNING: python-docx not installed — skipping DOCX generation", file=sys.stderr)
        return None

    doc = docx.Document()

    # Page setup — US Letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_heading("Salesforce Org Audit Report", level=0)
    for run in title.runs:
        run.font.color.rgb = _hex_to_rgb(BRAND_BLUE)

    doc.add_paragraph(
        f"{org_name} | Org ID: {org_id} | Instance: {instance} | {run_date}"
    )

    overall = summary["overall_score"]
    rating = summary["overall_rating"]
    p = doc.add_paragraph()
    run = p.add_run(f"Overall Score: {overall:.0f}/100 — {rating}")
    run.bold = True
    run.font.size = Pt(14)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    counts = summary["counts"]
    inv_map = [
        ("apex_classes", "Apex Classes"),
        ("apex_triggers", "Apex Triggers"),
        ("active_flows", "Active Flows"),
        ("process_builders", "Process Builders"),
        ("lwc_bundles", "LWC Components"),
        ("custom_objects", "Custom Objects"),
        ("validation_rules", "Validation Rules"),
        ("formula_fields", "Formula Fields"),
        ("workflow_rules", "Workflow Rules"),
        ("approval_processes", "Approval Processes"),
        ("permission_sets", "Permission Sets"),
        ("permission_set_groups", "Permission Set Groups"),
        ("profiles", "Profiles"),
    ]
    for key, label in inv_map:
        doc.add_paragraph(f"{counts.get(key, 0)} {label}", style="List Bullet")

    # Domain Scores
    doc.add_heading("Domain Scores", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Domain", "Score %", "Rating", "Below 70%"]):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True

    for domain, label in [("apex", "Apex"), ("flows", "Flows"), ("lwc", "LWC"), ("metadata", "Metadata")]:
        pct = summary["domain_scores"].get(domain, 0)
        r = score_rating(pct)[0]
        below = summary["below_threshold"].get(domain, 0)
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{pct}%"
        row[2].text = r
        row[3].text = str(below)

    # Scored sections
    for section_title, items_key, max_label in [
        ("Apex Classes", "apex_scores", 150),
        ("Flows", "flow_scores", 110),
        ("LWC Components", "lwc_scores", 165),
        ("Data Model", "metadata_scores", 120),
    ]:
        items = data.get(items_key, [])
        if not items:
            continue
        doc.add_heading(section_title, level=1)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Score", "Top Issues"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(items, key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = f"{item.get('score', 0)}/{item.get('max_score', max_label)}"
            row[2].text = issues_str

    # Apex Triggers
    if data.get("trigger_findings"):
        doc.add_heading("Apex Triggers", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Object", "Events", "Severity", "Findings"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(data["trigger_findings"], key=lambda x: x.get("name", "")):
            findings = item.get("findings", [])
            findings_str = "; ".join(
                f.get("message", f.get("finding", "")) for f in findings[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if findings else "LOW"
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = item.get("object", "")
            row[2].text = item.get("events", "")
            row[3].text = sev
            row[4].text = findings_str

    # Process Builders
    if data.get("process_builders"):
        doc.add_heading("Process Builders", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Object", "Criteria", "Actions", "Migration Priority"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(data["process_builders"], key=lambda x: x.get("name", "")):
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = item.get("object", "")
            row[2].text = str(item.get("criteria_count", ""))
            row[3].text = item.get("actions_summary", "")
            row[4].text = item.get("migration_priority", "HIGH")

    # Permission Findings
    if data.get("permission_findings"):
        doc.add_heading("Profiles & Permissions", level=1)
        for f in sorted(
            data["permission_findings"],
            key=lambda x: (
                {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}.get(
                    x.get("severity", "LOW").upper(), 4
                ),
                x.get("name", ""),
            ),
        ):
            doc.add_paragraph(
                f"[{f.get('severity', 'MEDIUM')}] {f.get('message', f.get('finding', ''))}",
                style="List Bullet",
            )

    # Unused Fields & Objects
    if data.get("unused_fields") or data.get("unused_objects"):
        doc.add_heading("Unused Fields & Objects", level=1)
        if data.get("unused_fields"):
            doc.add_heading("Unused Fields", level=2)
            table = doc.add_table(rows=1, cols=7)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, h in enumerate(["Object", "Field", "Data Type", "Has Data", "Referenced In", "Category", "Severity"]):
                hdr[i].text = h
                for run in hdr[i].paragraphs[0].runs:
                    run.font.bold = True
            for item in sorted(data["unused_fields"], key=lambda x: (
                {"HIGH": 0, "MEDIUM": 1, "LOW": 2}.get(x.get("severity", "LOW").upper(), 3),
                x.get("object", ""), x.get("field", ""),
            )):
                has_data = item.get("has_data")
                row = table.add_row().cells
                row[0].text = item.get("object", "")
                row[1].text = item.get("field", "")
                row[2].text = item.get("data_type", "")
                row[3].text = "Unknown" if has_data is None else ("Yes" if has_data else "No")
                row[4].text = ", ".join(item.get("referenced_in") or []) or "None"
                row[5].text = item.get("category", "")
                row[6].text = item.get("severity", "LOW")
        if data.get("unused_objects"):
            doc.add_heading("Unused Objects", level=2)
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, h in enumerate(["Object", "Record Count", "Referenced In", "Category", "Severity"]):
                hdr[i].text = h
                for run in hdr[i].paragraphs[0].runs:
                    run.font.bold = True
            for item in sorted(data["unused_objects"], key=lambda x: (
                {"HIGH": 0, "MEDIUM": 1, "LOW": 2}.get(x.get("severity", "LOW").upper(), 3),
                x.get("object", ""),
            )):
                row = table.add_row().cells
                row[0].text = item.get("object", "")
                row[1].text = str(item.get("record_count", 0))
                row[2].text = ", ".join(item.get("referenced_in") or []) or "None"
                row[3].text = item.get("category", "")
                row[4].text = item.get("severity", "LOW")

    # Validation Rules
    if data.get("validation_rules"):
        doc.add_heading("Validation Rules", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Object", "Active", "Severity", "Findings"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(data["validation_rules"], key=lambda x: x.get("name", "")):
            findings = item.get("findings", [])
            findings_str = "; ".join(
                f.get("message", f.get("finding", "")) for f in findings[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if findings else "LOW"
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = item.get("object", "")
            row[2].text = str(item.get("active", ""))
            row[3].text = sev
            row[4].text = findings_str

    # Formula Fields
    if data.get("formula_fields"):
        doc.add_heading("Formula Fields", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Object", "Data Type", "Length", "Severity", "Findings"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(data["formula_fields"], key=lambda x: x.get("name", "")):
            findings = item.get("findings", [])
            findings_str = "; ".join(
                f.get("message", "") for f in findings[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if findings else "—"
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = item.get("object", "")
            row[2].text = item.get("data_type", "")
            row[3].text = str(item.get("formula_length", ""))
            row[4].text = sev
            row[5].text = findings_str

    # Workflow Rules
    if data.get("workflow_rules"):
        doc.add_heading("Workflow Rules", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Object", "Action Types", "Priority", "Formula Findings"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(data["workflow_rules"], key=lambda x: x.get("name", "")):
            findings_str = "; ".join(
                f.get("message", "") for f in item.get("findings", [])[:3]
            )
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = item.get("object", "")
            row[2].text = item.get("action_types", "")
            row[3].text = item.get("migration_priority", "HIGH")
            row[4].text = findings_str

    # Other Declarative Logic
    if data.get("other_rules_findings"):
        doc.add_heading("Other Declarative Logic", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Type", "Name", "Object", "Severity", "Findings"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(data["other_rules_findings"], key=lambda x: x.get("name", "")):
            findings = item.get("findings", [])
            findings_str = "; ".join(
                f.get("message", "") for f in findings[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if findings else "—"
            row = table.add_row().cells
            row[0].text = item.get("type", "")
            row[1].text = item.get("name", "")
            row[2].text = item.get("object", "")
            row[3].text = sev
            row[4].text = findings_str

    # Hardcoded Values Summary
    doc.add_heading("Hardcoded Values Summary", level=1)
    hv_rows = _collect_hardcoded_values(data)
    if hv_rows:
        doc.add_paragraph(
            f"{len(hv_rows)} hardcoded value(s) found across all declarative logic."
        )
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Component Type", "Name", "Severity", "Finding"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for hv in hv_rows:
            row = table.add_row().cells
            row[0].text = hv["component_type"]
            row[1].text = hv["name"]
            row[2].text = hv["severity"]
            row[3].text = hv["message"]
    else:
        doc.add_paragraph("No hardcoded values detected in formulas or declarative logic.")

    # Reports & Dashboards
    rd = data.get("reports_dashboards", {})
    rd_items = rd.get("reports", []) + rd.get("dashboards", [])
    if rd_items:
        doc.add_heading("Reports & Dashboards", level=1)
        stale = sum(1 for r in rd_items if r.get("is_stale"))
        doc.add_paragraph(f"{len(rd_items)} total items, {stale} stale")
        for f in rd.get("findings", []):
            doc.add_paragraph(
                f"[{f.get('severity', 'MEDIUM')}] {f.get('message', '')}",
                style="List Bullet",
            )

    # Integrations
    if data.get("integrations"):
        doc.add_heading("Integration Analysis", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Type", "Name", "Endpoint", "Findings"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in data["integrations"]:
            findings_str = "; ".join(
                f.get("message", "") for f in item.get("findings", [])[:3]
            )
            row = table.add_row().cells
            row[0].text = item.get("type", "")
            row[1].text = item.get("name", "")
            row[2].text = item.get("endpoint", "") or "—"
            row[3].text = findings_str

    # Test Coverage
    tc = data.get("test_coverage", {})
    tc_classes = tc.get("classes", [])
    if tc_classes:
        doc.add_heading("Test Coverage", level=1)
        doc.add_paragraph(f"Org-wide coverage: {tc.get('org_wide_pct') or 0.0:.1f}%")
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Class/Trigger", "Covered", "Uncovered", "Coverage %"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(tc_classes, key=lambda x: x.get("coverage_pct", 0)):
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = str(item.get("lines_covered", 0))
            row[2].text = str(item.get("lines_uncovered", 0))
            row[3].text = f'{item.get("coverage_pct", 0):.1f}%'

    # Licensing
    lic = data.get("licensing", {})
    lic_all = lic.get("user_licenses", []) + lic.get("permission_set_licenses", []) + lic.get("package_licenses", [])
    if lic_all:
        doc.add_heading("Licensing Analysis", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Total", "Used", "Available", "Utilization"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in lic.get("user_licenses", []):
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = str(item.get("total", ""))
            row[2].text = str(item.get("used", ""))
            row[3].text = str(item.get("available", ""))
            row[4].text = f'{item.get("utilization_pct", 0):.0f}%'
        for item in lic.get("permission_set_licenses", []):
            row = table.add_row().cells
            row[0].text = item.get("label", item.get("name", ""))
            row[1].text = str(item.get("total", ""))
            row[2].text = str(item.get("used", ""))
            row[3].text = str(item.get("available", ""))
            row[4].text = f'{item.get("utilization_pct", 0):.0f}%'
        for item in lic.get("package_licenses", []):
            allowed = item.get("allowed", 0)
            used = item.get("used", 0)
            row = table.add_row().cells
            row[0].text = item.get("namespace", "")
            row[1].text = str(allowed)
            row[2].text = str(used)
            row[3].text = str(allowed - used)
            row[4].text = f'{(used / allowed * 100) if allowed > 0 else 0:.0f}%'
        for f in lic.get("findings", []):
            doc.add_paragraph(
                f"[{f.get('severity', 'MEDIUM')}] {f.get('message', '')}",
                style="List Bullet",
            )

    # Team Evaluation
    team = data.get("team_evaluation", {})
    if team.get("users"):
        doc.add_heading("Team Evaluation", level=1)
        doc.add_paragraph(f"{team.get('active_users', 0)} active users")
        for f in team.get("findings", []):
            doc.add_paragraph(
                f"[{f.get('severity', 'MEDIUM')}] {f.get('message', '')}",
                style="List Bullet",
            )

    # Change History
    ch = data.get("change_history", {})
    if ch.get("deployments") or ch.get("audit_trail"):
        doc.add_heading("Change History", level=1)
        doc.add_paragraph(
            f"Deployment success rate: {ch.get('deployment_success_rate', 0):.0f}% "
            f"({len(ch.get('deployments', []))} deployments)"
        )
        for f in ch.get("findings", []):
            doc.add_paragraph(
                f"[{f.get('severity', 'MEDIUM')}] {f.get('message', '')}",
                style="List Bullet",
            )

    # Recommendations
    doc.add_heading("Recommendations", level=1)
    rec_data = _recommendations_list(data)
    if rec_data:
        for rec in rec_data:
            doc.add_paragraph(rec, style="List Number")
    else:
        doc.add_paragraph("No recommendations — org is in great shape.")

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Generated by Salesforce Core Skills Audit Engine | {run_date} | Org: {org_name} ({org_id})")
    run.font.size = Pt(9)
    run.font.color.rgb = _hex_to_rgb(MUTED)

    doc.save(output_path)
    return output_path


# ── XLSX report ─────────────────────────────────────────────────────────────


def generate_xlsx(data, summary, org_name, org_id, instance, run_date, output_path):
    """Generate the Excel audit workbook."""
    if not HAS_OPENPYXL:
        print("WARNING: openpyxl not installed — skipping XLSX generation", file=sys.stderr)
        return None

    wb = openpyxl.Workbook()

    header_fill = PatternFill("solid", fgColor="417AE4")
    header_font = Font(bold=True, color="FFFFFF", name="Arial", size=11)

    def apply_header(ws, headers):
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="left", vertical="center")

    def auto_width(ws):
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    score_fills = {
        "Excellent": PatternFill("solid", fgColor="E8FBF9"),
        "Good": PatternFill("solid", fgColor="E9F7EF"),
        "Acceptable": PatternFill("solid", fgColor="FEF3CD"),
        "Warning": PatternFill("solid", fgColor="FEF9E7"),
        "Critical": PatternFill("solid", fgColor="FDE8E8"),
    }
    score_fonts = {
        "Excellent": Font(color="14DDDD", bold=True, name="Arial", size=11),
        "Good": Font(color="27AE60", bold=True, name="Arial", size=11),
        "Acceptable": Font(color="F39C12", bold=True, name="Arial", size=11),
        "Warning": Font(color="E67E22", bold=True, name="Arial", size=11),
        "Critical": Font(color="E74C3C", bold=True, name="Arial", size=11),
    }

    def style_score_cell(cell, score, max_score):
        pct = (score / max_score * 100) if max_score > 0 else 0
        rating = score_rating(pct)[0]
        # Map "Needs Improvement" to "Warning" for styling
        style_key = rating if rating in score_fills else "Warning"
        cell.fill = score_fills.get(style_key, score_fills["Warning"])
        cell.font = score_fonts.get(style_key, score_fonts["Warning"])

    # Sheet 1 — Apex Classes
    ws = wb.active
    ws.title = "Apex Classes"
    ws.sheet_properties.tabColor = "417AE4"
    apply_header(ws, ["Name", "Score", "Max", "Rating", "Top Issues"])
    for i, item in enumerate(sorted(data.get("apex_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 150)
        pct = (s / m * 100) if m > 0 else 0
        ws.cell(row=i, column=1, value=item.get("name", ""))
        sc = ws.cell(row=i, column=2, value=s)
        style_score_cell(sc, s, m)
        ws.cell(row=i, column=3, value=m)
        ws.cell(row=i, column=4, value=score_rating(pct)[0])
        issues = item.get("issues", [])
        ws.cell(row=i, column=5, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in issues[:3]
        ))
    auto_width(ws)

    # Sheet 2 — Apex Triggers
    ws2 = wb.create_sheet("Apex Triggers")
    ws2.sheet_properties.tabColor = "417AE4"
    apply_header(ws2, ["Name", "Object", "Events", "Findings", "Severity"])
    for i, item in enumerate(data.get("trigger_findings", []), 2):
        ws2.cell(row=i, column=1, value=item.get("name", ""))
        ws2.cell(row=i, column=2, value=item.get("object", ""))
        ws2.cell(row=i, column=3, value=item.get("events", ""))
        findings = item.get("findings", [])
        ws2.cell(row=i, column=4, value="; ".join(
            f.get("message", f.get("finding", "")) for f in findings[:3]
        ))
        if findings:
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
            )
            ws2.cell(row=i, column=5, value=sev)
    auto_width(ws2)

    # Sheet 3 — Flows
    ws3 = wb.create_sheet("Flows")
    ws3.sheet_properties.tabColor = "417AE4"
    apply_header(ws3, ["Name", "Process Type", "Score", "Max", "Top Issues"])
    for i, item in enumerate(sorted(data.get("flow_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 110)
        ws3.cell(row=i, column=1, value=item.get("name", ""))
        ws3.cell(row=i, column=2, value=item.get("process_type", ""))
        sc = ws3.cell(row=i, column=3, value=s)
        style_score_cell(sc, s, m)
        ws3.cell(row=i, column=4, value=m)
        ws3.cell(row=i, column=5, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in item.get("issues", [])[:3]
        ))
    auto_width(ws3)

    # Sheet 4 — Process Builders
    ws4 = wb.create_sheet("Process Builders")
    ws4.sheet_properties.tabColor = "417AE4"
    apply_header(ws4, ["Name", "Object", "Criteria Count", "Actions", "Migration Priority"])
    for i, item in enumerate(data.get("process_builders", []), 2):
        ws4.cell(row=i, column=1, value=item.get("name", ""))
        ws4.cell(row=i, column=2, value=item.get("object", ""))
        ws4.cell(row=i, column=3, value=item.get("criteria_count", 0))
        ws4.cell(row=i, column=4, value=item.get("actions_summary", ""))
        ws4.cell(row=i, column=5, value=item.get("migration_priority", "HIGH"))
    auto_width(ws4)

    # Sheet 5 — LWC
    ws5 = wb.create_sheet("LWC")
    ws5.sheet_properties.tabColor = "417AE4"
    apply_header(ws5, ["Name", "Score", "Max", "Rating", "Top Issues"])
    for i, item in enumerate(sorted(data.get("lwc_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 165)
        pct = (s / m * 100) if m > 0 else 0
        ws5.cell(row=i, column=1, value=item.get("name", ""))
        sc = ws5.cell(row=i, column=2, value=s)
        style_score_cell(sc, s, m)
        ws5.cell(row=i, column=3, value=m)
        ws5.cell(row=i, column=4, value=score_rating(pct)[0])
        ws5.cell(row=i, column=5, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in item.get("issues", [])[:3]
        ))
    auto_width(ws5)

    # Sheet 6 — Profiles
    ws6 = wb.create_sheet("Profiles")
    ws6.sheet_properties.tabColor = "417AE4"
    apply_header(ws6, ["Name", "UserType", "Key Permissions", "Findings", "Severity"])
    # Profiles come from permission_findings filtered
    profile_findings = [f for f in data.get("permission_findings", []) if f.get("type") == "profile"]
    for i, item in enumerate(profile_findings, 2):
        ws6.cell(row=i, column=1, value=item.get("name", ""))
        ws6.cell(row=i, column=2, value=item.get("user_type", ""))
        ws6.cell(row=i, column=3, value=item.get("key_permissions", ""))
        ws6.cell(row=i, column=4, value=item.get("message", item.get("finding", "")))
        ws6.cell(row=i, column=5, value=item.get("severity", ""))
    auto_width(ws6)

    # Sheet 7 — Permission Sets
    ws7 = wb.create_sheet("Permission Sets")
    ws7.sheet_properties.tabColor = "417AE4"
    apply_header(ws7, ["Name", "Label", "Assignments", "Findings", "Severity"])
    ps_findings = [f for f in data.get("permission_findings", []) if f.get("type") == "permission_set"]
    for i, item in enumerate(ps_findings, 2):
        ws7.cell(row=i, column=1, value=item.get("name", ""))
        ws7.cell(row=i, column=2, value=item.get("label", ""))
        ws7.cell(row=i, column=3, value=item.get("assignments", 0))
        ws7.cell(row=i, column=4, value=item.get("message", item.get("finding", "")))
        ws7.cell(row=i, column=5, value=item.get("severity", ""))
    auto_width(ws7)

    # Sheet 8 — Custom Objects
    ws8 = wb.create_sheet("Custom Objects")
    ws8.sheet_properties.tabColor = "417AE4"
    apply_header(ws8, ["Name", "Score", "Max", "Fields", "Relationships", "Top Issues"])
    for i, item in enumerate(sorted(data.get("metadata_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 120)
        ws8.cell(row=i, column=1, value=item.get("name", ""))
        sc = ws8.cell(row=i, column=2, value=s)
        style_score_cell(sc, s, m)
        ws8.cell(row=i, column=3, value=m)
        ws8.cell(row=i, column=4, value=item.get("field_count", 0))
        ws8.cell(row=i, column=5, value=item.get("relationship_count", 0))
        ws8.cell(row=i, column=6, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in item.get("issues", [])[:3]
        ))
    auto_width(ws8)

    # Sheet 9 — Unused Fields
    ws9 = wb.create_sheet("Unused Fields")
    ws9.sheet_properties.tabColor = "417AE4"
    apply_header(ws9, ["Object", "Field", "Data Type", "Has Data", "Referenced In", "Category", "Severity"])
    for i, item in enumerate(sorted(
        data.get("unused_fields", []),
        key=lambda x: (
            {"HIGH": 0, "MEDIUM": 1, "LOW": 2}.get(x.get("severity", "LOW").upper(), 3),
            x.get("object", ""), x.get("field", ""),
        ),
    ), 2):
        has_data = item.get("has_data")
        ws9.cell(row=i, column=1, value=item.get("object", ""))
        ws9.cell(row=i, column=2, value=item.get("field", ""))
        ws9.cell(row=i, column=3, value=item.get("data_type", ""))
        ws9.cell(row=i, column=4, value="Unknown" if has_data is None else ("Yes" if has_data else "No"))
        ws9.cell(row=i, column=5, value=", ".join(item.get("referenced_in") or []) or "None")
        ws9.cell(row=i, column=6, value=item.get("category", ""))
        ws9.cell(row=i, column=7, value=item.get("severity", "LOW"))
    auto_width(ws9)

    # Sheet 10 — Unused Objects
    ws10 = wb.create_sheet("Unused Objects")
    ws10.sheet_properties.tabColor = "417AE4"
    apply_header(ws10, ["Object", "Record Count", "Referenced In", "Category", "Severity"])
    for i, item in enumerate(sorted(
        data.get("unused_objects", []),
        key=lambda x: (
            {"HIGH": 0, "MEDIUM": 1, "LOW": 2}.get(x.get("severity", "LOW").upper(), 3),
            x.get("object", ""),
        ),
    ), 2):
        ws10.cell(row=i, column=1, value=item.get("object", ""))
        ws10.cell(row=i, column=2, value=item.get("record_count", 0))
        ws10.cell(row=i, column=3, value=", ".join(item.get("referenced_in") or []) or "None")
        ws10.cell(row=i, column=4, value=item.get("category", ""))
        ws10.cell(row=i, column=5, value=item.get("severity", "LOW"))
    auto_width(ws10)

    # Sheet 11 — Validation Rules
    ws11 = wb.create_sheet("Validation Rules")
    ws11.sheet_properties.tabColor = "417AE4"
    apply_header(ws11, ["Name", "Object", "Active", "Findings", "Severity"])
    for i, item in enumerate(data.get("validation_rules", []), 2):
        ws11.cell(row=i, column=1, value=item.get("name", ""))
        ws11.cell(row=i, column=2, value=item.get("object", ""))
        ws11.cell(row=i, column=3, value=str(item.get("active", "")))
        findings = item.get("findings", [])
        ws11.cell(row=i, column=4, value="; ".join(
            f.get("message", f.get("finding", "")) for f in findings[:3]
        ))
        if findings:
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
            )
            ws11.cell(row=i, column=5, value=sev)
    auto_width(ws11)

    # Sheet 12 — Formula Fields
    ws12 = wb.create_sheet("Formula Fields")
    ws12.sheet_properties.tabColor = "417AE4"
    apply_header(ws12, ["Name", "Object", "Data Type", "Formula Length", "Findings", "Severity"])
    for i, item in enumerate(data.get("formula_fields", []), 2):
        ws12.cell(row=i, column=1, value=item.get("name", ""))
        ws12.cell(row=i, column=2, value=item.get("object", ""))
        ws12.cell(row=i, column=3, value=item.get("data_type", ""))
        ws12.cell(row=i, column=4, value=item.get("formula_length", 0))
        findings = item.get("findings", [])
        ws12.cell(row=i, column=5, value="; ".join(
            f.get("message", "") for f in findings[:3]
        ))
        if findings:
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
            )
            ws12.cell(row=i, column=6, value=sev)
    auto_width(ws12)

    # Sheet 13 — Workflow Rules
    ws13 = wb.create_sheet("Workflow Rules")
    ws13.sheet_properties.tabColor = "417AE4"
    apply_header(ws13, ["Name", "Object", "Action Types", "Migration Priority", "Formula Findings"])
    for i, item in enumerate(data.get("workflow_rules", []), 2):
        ws13.cell(row=i, column=1, value=item.get("name", ""))
        ws13.cell(row=i, column=2, value=item.get("object", ""))
        ws13.cell(row=i, column=3, value=item.get("action_types", ""))
        ws13.cell(row=i, column=4, value=item.get("migration_priority", "HIGH"))
        findings = item.get("findings", [])
        ws13.cell(row=i, column=5, value="; ".join(
            f.get("message", "") for f in findings[:3]
        ))
    auto_width(ws13)

    # Sheet 14 — Other Declarative Logic
    ws14 = wb.create_sheet("Other Declarative Logic")
    ws14.sheet_properties.tabColor = "417AE4"
    apply_header(ws14, ["Type", "Name", "Object", "Findings", "Severity"])
    for i, item in enumerate(data.get("other_rules_findings", []), 2):
        ws14.cell(row=i, column=1, value=item.get("type", ""))
        ws14.cell(row=i, column=2, value=item.get("name", ""))
        ws14.cell(row=i, column=3, value=item.get("object", ""))
        findings = item.get("findings", [])
        ws14.cell(row=i, column=4, value="; ".join(
            f.get("message", "") for f in findings[:3]
        ))
        if findings:
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
            )
            ws14.cell(row=i, column=5, value=sev)
    auto_width(ws14)

    # Sheet 15 — Hardcoded Values
    ws15 = wb.create_sheet("Hardcoded Values")
    ws15.sheet_properties.tabColor = "417AE4"
    apply_header(ws15, ["Component Type", "Name", "Severity", "Finding"])
    hv_rows = _collect_hardcoded_values(data)
    for i, hv in enumerate(hv_rows, 2):
        ws15.cell(row=i, column=1, value=hv["component_type"])
        ws15.cell(row=i, column=2, value=hv["name"])
        ws15.cell(row=i, column=3, value=hv["severity"])
        ws15.cell(row=i, column=4, value=hv["message"])
    auto_width(ws15)

    # Sheet 16 — Reports & Dashboards
    ws16_rd = wb.create_sheet("Reports & Dashboards")
    ws16_rd.sheet_properties.tabColor = "417AE4"
    apply_header(ws16_rd, ["Name", "Type", "Folder", "Last Activity", "Created", "Stale"])
    rd = data.get("reports_dashboards", {})
    rd_row = 2
    for item in rd.get("reports", []):
        ws16_rd.cell(row=rd_row, column=1, value=item.get("name", ""))
        ws16_rd.cell(row=rd_row, column=2, value="Report")
        ws16_rd.cell(row=rd_row, column=3, value=item.get("folder", ""))
        ws16_rd.cell(row=rd_row, column=4, value=item.get("last_run_date") or "Never")
        ws16_rd.cell(row=rd_row, column=5, value=item.get("created_date", ""))
        ws16_rd.cell(row=rd_row, column=6, value="Yes" if item.get("is_stale") else "No")
        rd_row += 1
    for item in rd.get("dashboards", []):
        ws16_rd.cell(row=rd_row, column=1, value=item.get("name", ""))
        ws16_rd.cell(row=rd_row, column=2, value="Dashboard")
        ws16_rd.cell(row=rd_row, column=3, value=item.get("folder", ""))
        ws16_rd.cell(row=rd_row, column=4, value=item.get("last_viewed_date") or "Never")
        ws16_rd.cell(row=rd_row, column=5, value=item.get("created_date", ""))
        ws16_rd.cell(row=rd_row, column=6, value="Yes" if item.get("is_stale") else "No")
        rd_row += 1
    auto_width(ws16_rd)

    # Sheet 17 — Integrations
    ws17_int = wb.create_sheet("Integrations")
    ws17_int.sheet_properties.tabColor = "417AE4"
    apply_header(ws17_int, ["Type", "Name", "Endpoint", "Findings", "Severity"])
    for i, item in enumerate(data.get("integrations", []), 2):
        ws17_int.cell(row=i, column=1, value=item.get("type", ""))
        ws17_int.cell(row=i, column=2, value=item.get("name", ""))
        ws17_int.cell(row=i, column=3, value=item.get("endpoint", "") or "")
        findings = item.get("findings", [])
        ws17_int.cell(row=i, column=4, value="; ".join(
            f.get("message", "") for f in findings[:3]
        ))
        if findings:
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
            )
            ws17_int.cell(row=i, column=5, value=sev)
    auto_width(ws17_int)

    # Sheet 18 — Test Coverage
    ws18_tc = wb.create_sheet("Test Coverage")
    ws18_tc.sheet_properties.tabColor = "417AE4"
    apply_header(ws18_tc, ["Class/Trigger", "Lines Covered", "Lines Uncovered", "Coverage %"])
    tc = data.get("test_coverage", {})
    for i, item in enumerate(sorted(tc.get("classes", []), key=lambda x: x.get("coverage_pct", 0)), 2):
        ws18_tc.cell(row=i, column=1, value=item.get("name", ""))
        ws18_tc.cell(row=i, column=2, value=item.get("lines_covered", 0))
        ws18_tc.cell(row=i, column=3, value=item.get("lines_uncovered", 0))
        pct = item.get("coverage_pct", 0)
        c = ws18_tc.cell(row=i, column=4, value=round(pct, 1))
        if pct < 75:
            c.font = Font(color="E74C3C", bold=True, name="Arial", size=11)
    auto_width(ws18_tc)

    # Sheet 19 — Team
    ws19_team = wb.create_sheet("Team")
    ws19_team.sheet_properties.tabColor = "417AE4"
    apply_header(ws19_team, ["Name", "Username", "Profile", "Role", "Last Login", "Days Inactive", "PS Count"])
    team = data.get("team_evaluation", {})
    for i, item in enumerate(sorted(team.get("users", []),
                                    key=lambda x: x.get("days_since_login") or 0, reverse=True), 2):
        ws19_team.cell(row=i, column=1, value=item.get("name", ""))
        ws19_team.cell(row=i, column=2, value=item.get("username", ""))
        ws19_team.cell(row=i, column=3, value=item.get("profile", ""))
        ws19_team.cell(row=i, column=4, value=item.get("role", "") or "")
        ws19_team.cell(row=i, column=5, value=item.get("last_login", ""))
        ws19_team.cell(row=i, column=6, value=item.get("days_since_login") or "")
        ws19_team.cell(row=i, column=7, value=item.get("permission_set_count", 0))
    auto_width(ws19_team)

    # Sheet 20 — Change History
    ws20_ch = wb.create_sheet("Change History")
    ws20_ch.sheet_properties.tabColor = "417AE4"
    apply_header(ws20_ch, ["Date", "User", "Status", "Components Deployed", "Errors"])
    ch = data.get("change_history", {})
    for i, item in enumerate(ch.get("deployments", []), 2):
        ws20_ch.cell(row=i, column=1, value=item.get("date", ""))
        ws20_ch.cell(row=i, column=2, value=item.get("user", ""))
        ws20_ch.cell(row=i, column=3, value=item.get("status", ""))
        ws20_ch.cell(row=i, column=4, value=item.get("components_deployed", 0))
        ws20_ch.cell(row=i, column=5, value=item.get("component_errors", 0))
    auto_width(ws20_ch)

    # Sheet 21 — Licensing
    ws21_lic = wb.create_sheet("Licensing")
    ws21_lic.sheet_properties.tabColor = "417AE4"
    apply_header(ws21_lic, ["Name", "Type", "Total", "Used", "Available", "Utilization %"])
    lic = data.get("licensing", {})
    lic_row = 2
    for item in lic.get("user_licenses", []):
        ws21_lic.cell(row=lic_row, column=1, value=item.get("name", ""))
        ws21_lic.cell(row=lic_row, column=2, value="User License")
        ws21_lic.cell(row=lic_row, column=3, value=item.get("total", 0))
        ws21_lic.cell(row=lic_row, column=4, value=item.get("used", 0))
        ws21_lic.cell(row=lic_row, column=5, value=item.get("available", 0))
        ws21_lic.cell(row=lic_row, column=6, value=round(item.get("utilization_pct", 0), 1))
        lic_row += 1
    for item in lic.get("permission_set_licenses", []):
        ws21_lic.cell(row=lic_row, column=1, value=item.get("label", item.get("name", "")))
        ws21_lic.cell(row=lic_row, column=2, value="PS License")
        ws21_lic.cell(row=lic_row, column=3, value=item.get("total", 0))
        ws21_lic.cell(row=lic_row, column=4, value=item.get("used", 0))
        ws21_lic.cell(row=lic_row, column=5, value=item.get("available", 0))
        ws21_lic.cell(row=lic_row, column=6, value=round(item.get("utilization_pct", 0), 1))
        lic_row += 1
    for item in lic.get("package_licenses", []):
        allowed = item.get("allowed", 0)
        used = item.get("used", 0)
        ws21_lic.cell(row=lic_row, column=1, value=item.get("namespace", ""))
        ws21_lic.cell(row=lic_row, column=2, value="Package License")
        ws21_lic.cell(row=lic_row, column=3, value=allowed)
        ws21_lic.cell(row=lic_row, column=4, value=used)
        ws21_lic.cell(row=lic_row, column=5, value=allowed - used)
        ws21_lic.cell(row=lic_row, column=6, value=round((used / allowed * 100) if allowed > 0 else 0, 1))
        lic_row += 1
    auto_width(ws21_lic)

    # Sheet 22 — Summary
    ws22 = wb.create_sheet("Summary")
    ws22.sheet_properties.tabColor = "417AE4"
    apply_header(ws22, ["Metric", "Value"])
    summary_rows = [
        ("Org Name", org_name),
        ("Org ID", org_id),
        ("Instance", instance),
        ("Run Date", run_date),
        ("Overall Score", f"{summary['overall_score']:.0f}/100"),
        ("Overall Rating", summary["overall_rating"]),
        ("", ""),
        ("Domain Scores", ""),
    ]
    for domain in ["apex", "flows", "lwc", "metadata"]:
        pct = summary["domain_scores"].get(domain, 0)
        summary_rows.append((f"  {domain.title()}", f"{pct}%"))

    summary_rows.append(("", ""))
    summary_rows.append(("Severity Counts", ""))
    for sev in ["CRITICAL", "HIGH", "MEDIUM", "LOW"]:
        summary_rows.append((f"  {sev}", summary["severity_counts"].get(sev, 0)))

    for i, (metric, value) in enumerate(summary_rows, 2):
        ws22.cell(row=i, column=1, value=metric)
        ws22.cell(row=i, column=2, value=value)
    auto_width(ws22)

    wb.save(output_path)
    return output_path


# ── JSON summary ────────────────────────────────────────────────────────────


def generate_json_summary(summary, run_date, output_path):
    """Generate the JSON audit summary."""
    output = {**summary, "generated_date": run_date}
    Path(output_path).write_text(json.dumps(output, indent=2, default=str), encoding="utf-8")
    return output_path


# ── Standalone report generators ──────────────────────────────────────────


def _standalone_html(title, org_name, run_date, body_html):
    """Wrap body HTML in a branded standalone page."""
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{_esc(title)} — {_esc(org_name)}</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
       background: {BODY_BG}; color: #1a1a2e; line-height: 1.5; }}
.banner {{ background: {BRAND_BLUE}; padding: 24px 40px; }}
.banner-title {{ font-size: 22px; font-weight: 700; color: #fff; }}
.banner-subtitle {{ font-size: 12px; color: rgba(255,255,255,.80); margin-top: 4px; }}
.card {{ background: #fff; border: 1px solid {BORDER}; border-radius: 8px;
         padding: 24px; margin: 16px 40px; }}
.card h2 {{ color: {BRAND_BLUE}; font-size: 18px; margin-bottom: 16px; }}
table {{ width: 100%; border-collapse: collapse; margin-top: 8px; }}
thead th {{ background: {BRAND_BLUE}; color: #fff; font-size: 11px; font-weight: 700;
  text-transform: uppercase; letter-spacing: 0.6px; padding: 10px 14px; text-align: left; }}
tbody td {{ padding: 10px 14px; border-bottom: 1px solid {BORDER}; font-size: 13px; }}
tbody tr:hover {{ background: #f8f9fb; }}
.finding {{ border-left: 4px solid; border-radius: 6px; padding: 12px 16px; margin-bottom: 8px;
            background: #fff; }}
.finding.critical {{ border-left-color: #E74C3C; }}
.finding.warning {{ border-left-color: #E67E22; }}
.finding.info {{ border-left-color: {BRAND_BLUE}; }}
.finding.positive {{ border-left-color: #27AE60; }}
.finding-badge {{ display: inline-block; padding: 2px 8px; border-radius: 10px;
  font-size: 11px; font-weight: 600; margin-right: 8px; }}
.finding.critical .finding-badge {{ background: #FDE8E8; color: #E74C3C; }}
.finding.warning .finding-badge {{ background: #FEF3CD; color: #E67E22; }}
.finding.info .finding-badge {{ background: #EBF1FB; color: {BRAND_BLUE}; }}
.finding.positive .finding-badge {{ background: #E9F7EF; color: #27AE60; }}
.footer {{ background: {BRAND_BLUE}; padding: 16px 40px; text-align: center;
  font-size: 12px; color: rgba(255,255,255,.80); margin-top: 24px; }}
.footer a {{ color: #fff; }}
</style>
</head>
<body>
<div class="banner">
  <div class="banner-title">{_esc(title)}</div>
  <div class="banner-subtitle">{_esc(org_name)} &middot; {_esc(run_date)}</div>
</div>
{body_html}
<div class="footer">
  Generated by Salesforce Core Skills &middot; {_esc(run_date)}
</div>
</body>
</html>"""


def generate_standalone_reports(data, summary, org_name, run_date, output_dir):
    """Generate all 12 standalone report documents.

    Returns a list of generated file paths.
    """
    out = Path(output_dir)
    generated = []

    # ── 1. Reports & Dashboards Inventory (XLSX) ──
    if HAS_OPENPYXL:
        rd = data.get("reports_dashboards", {})
        if rd.get("reports") or rd.get("dashboards"):
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Reports"
            hf = PatternFill("solid", fgColor="417AE4")
            hfont = Font(bold=True, color="FFFFFF", name="Arial", size=11)
            for col, h in enumerate(["Name", "Folder", "Format", "Last Run", "Created", "Stale"], 1):
                c = ws.cell(row=1, column=col, value=h)
                c.fill = hf
                c.font = hfont
            for i, item in enumerate(rd.get("reports", []), 2):
                ws.cell(row=i, column=1, value=item.get("name", ""))
                ws.cell(row=i, column=2, value=item.get("folder", ""))
                ws.cell(row=i, column=3, value=item.get("format", ""))
                ws.cell(row=i, column=4, value=item.get("last_run_date") or "Never")
                ws.cell(row=i, column=5, value=item.get("created_date", ""))
                ws.cell(row=i, column=6, value="Yes" if item.get("is_stale") else "No")

            ws2 = wb.create_sheet("Dashboards")
            for col, h in enumerate(["Name", "Folder", "Last Viewed", "Created", "Stale"], 1):
                c = ws2.cell(row=1, column=col, value=h)
                c.fill = hf
                c.font = hfont
            for i, item in enumerate(rd.get("dashboards", []), 2):
                ws2.cell(row=i, column=1, value=item.get("name", ""))
                ws2.cell(row=i, column=2, value=item.get("folder", ""))
                ws2.cell(row=i, column=3, value=item.get("last_viewed_date") or "Never")
                ws2.cell(row=i, column=4, value=item.get("created_date", ""))
                ws2.cell(row=i, column=5, value="Yes" if item.get("is_stale") else "No")

            p = out / "Reports_Dashboards_Inventory.xlsx"
            wb.save(p)
            generated.append(str(p))

    # ── 2. Integration Analysis (HTML) ──
    items = data.get("integrations", [])
    if items:
        rows = []
        for item in items:
            findings_str = "; ".join(f.get("message", "") for f in item.get("findings", [])[:3])
            rows.append([
                _esc(item.get("type", "")), _esc(item.get("name", "")),
                _esc(item.get("endpoint", "") or "—"), _esc(findings_str),
            ])
        body = (
            f'<div class="card"><h2>Integrations ({len(items)})</h2>'
            f'{_table_html(["Type", "Name", "Endpoint", "Findings"], rows)}</div>'
        )
        p = out / "Integration_Analysis.html"
        p.write_text(_standalone_html("Integration Analysis", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 3. Test Coverage Report (HTML) ──
    tc = data.get("test_coverage", {})
    if tc.get("classes") or tc.get("org_wide_pct") is not None:
        rows = []
        for item in sorted(tc.get("classes", []), key=lambda x: x.get("coverage_pct", 0)):
            rows.append([
                _esc(item.get("name", "")), str(item.get("lines_covered", 0)),
                str(item.get("lines_uncovered", 0)), f'{item.get("coverage_pct", 0):.1f}%',
            ])
        body = (
            f'<div class="card"><h2>Org-Wide Coverage: {tc.get("org_wide_pct") or 0.0:.1f}%</h2>'
            f'{_table_html(["Class/Trigger", "Covered", "Uncovered", "Coverage %"], rows)}'
            f'{_findings_html(tc.get("findings", []))}</div>'
        )
        p = out / "Test_Coverage_Report.html"
        p.write_text(_standalone_html("Test Coverage Report", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 4. Licensing Analysis (HTML) ──
    lic = data.get("licensing", {})
    lic_all = lic.get("user_licenses", []) + lic.get("permission_set_licenses", []) + lic.get("package_licenses", [])
    if lic_all:
        rows = []
        for item in lic.get("user_licenses", []):
            rows.append([
                _esc(item.get("name", "")), "User License",
                str(item.get("total", "")), str(item.get("used", "")),
                f'{item.get("utilization_pct", 0):.0f}%',
            ])
        for item in lic.get("permission_set_licenses", []):
            rows.append([
                _esc(item.get("label", item.get("name", ""))), "PS License",
                str(item.get("total", "")), str(item.get("used", "")),
                f'{item.get("utilization_pct", 0):.0f}%',
            ])
        for item in lic.get("package_licenses", []):
            allowed = item.get("allowed", 0)
            used = item.get("used", 0)
            rows.append([
                _esc(item.get("namespace", "")), "Package License",
                str(allowed), str(used),
                f'{(used / allowed * 100) if allowed > 0 else 0:.0f}%',
            ])
        body = (
            f'<div class="card"><h2>License Utilisation</h2>'
            f'{_table_html(["Name", "Type", "Total", "Used", "Utilization"], rows)}'
            f'{_findings_html(lic.get("findings", []))}</div>'
        )
        p = out / "Licensing_Analysis.html"
        p.write_text(_standalone_html("Licensing Analysis", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 5. Team Evaluation (HTML) ──
    team = data.get("team_evaluation", {})
    if team.get("users"):
        rows = []
        for u in sorted(team["users"], key=lambda x: x.get("days_since_login") or 0, reverse=True):
            rows.append([
                _esc(u.get("name", "")), _esc(u.get("profile", "")),
                _esc(u.get("role", "") or "—"), _esc(u.get("last_login", "Never")),
                str(u.get("days_since_login", "—")),
            ])
        body = (
            f'<div class="card"><h2>Team Overview ({team.get("active_users", 0)} active users)</h2>'
            f'{_table_html(["Name", "Profile", "Role", "Last Login", "Days Inactive"], rows)}'
            f'{_findings_html(team.get("findings", []))}</div>'
        )
        p = out / "Team_Evaluation.html"
        p.write_text(_standalone_html("Team Evaluation", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 6. Change History Audit (HTML) ──
    ch = data.get("change_history", {})
    if ch.get("deployments") or ch.get("audit_trail"):
        dep_rows = []
        for item in ch.get("deployments", [])[:50]:
            dep_rows.append([
                _esc(item.get("date", "")), _esc(item.get("user", "")),
                _esc(item.get("status", "")), str(item.get("components_deployed", "")),
                str(item.get("component_errors", "")),
            ])
        trail_rows = []
        for item in ch.get("audit_trail", [])[:50]:
            trail_rows.append([
                _esc(item.get("date", "")), _esc(item.get("user", "")),
                _esc(item.get("action", "")), _esc(item.get("section", "")),
            ])
        body = (
            f'<div class="card"><h2>Deployments (success rate: '
            f'{ch.get("deployment_success_rate", 0):.0f}%)</h2>'
            f'{_table_html(["Date", "User", "Status", "Components", "Errors"], dep_rows)}</div>'
            f'<div class="card"><h2>Setup Audit Trail (recent)</h2>'
            f'{_table_html(["Date", "User", "Action", "Section"], trail_rows)}'
            f'{_findings_html(ch.get("findings", []))}</div>'
        )
        p = out / "Change_History_Audit.html"
        p.write_text(_standalone_html("Change History Audit", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 7. Data Quality Report (HTML) ──
    dq = data.get("data_quality", {})
    if dq.get("objects"):
        dq_rows = []
        for obj in dq["objects"]:
            for field in obj.get("field_completeness", []):
                dq_rows.append([
                    _esc(obj.get("name", "")),
                    str(obj.get("record_count", "")),
                    _esc(field.get("field", "")),
                    str(field.get("null_count", "")),
                    f'{field.get("null_pct", 0):.1f}%',
                ])
        body = (
            f'<div class="card"><h2>Field Completeness</h2>'
            f'{_table_html(["Object", "Records", "Field", "Null Count", "Null %"], dq_rows)}'
            f'{_findings_html(dq.get("findings", []))}</div>'
        )
        p = out / "Data_Quality_Report.html"
        p.write_text(_standalone_html("Data Quality Report", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 8. Technical Impact Assessment (HTML) ──
    # Synthesised from multiple domains
    impact_items = []
    # Low-scoring components
    for domain, key, max_s in [("Apex", "apex_scores", 150), ("Flows", "flow_scores", 110),
                                ("LWC", "lwc_scores", 165), ("Metadata", "metadata_scores", 120)]:
        for item in data.get(key, []):
            s = item.get("score", 0)
            m = item.get("max_score", max_s)
            pct = (s / m * 100) if m > 0 else 0
            if pct < 70:
                impact_items.append({
                    "component": item.get("name", ""),
                    "domain": domain,
                    "score_pct": round(pct, 1),
                    "risk": "HIGH" if pct < 40 else "MEDIUM",
                })
    # Zero-coverage classes
    for c in tc.get("classes", []):
        if c.get("coverage_pct", 0) == 0:
            impact_items.append({
                "component": c.get("name", ""),
                "domain": "Test Coverage",
                "score_pct": 0,
                "risk": "HIGH",
            })
    if impact_items:
        impact_items.sort(key=lambda x: x["score_pct"])
        rows = [[
            _esc(i["component"]), _esc(i["domain"]),
            f'{i["score_pct"]}%', _severity_badge_html(i["risk"]),
        ] for i in impact_items[:20]]
        body = (
            f'<div class="card"><h2>Highest Risk Components ({len(impact_items)} total)</h2>'
            f'{_table_html(["Component", "Domain", "Score", "Risk"], rows)}</div>'
        )
        p = out / "Technical_Impact_Assessment.html"
        p.write_text(_standalone_html("Technical Impact Assessment", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 9. Architectural Analysis (HTML) ──
    arch_sections = []
    # API version distribution
    api_versions = {}
    for item in data.get("apex_scores", []):
        v = str(item.get("api_version", "Unknown"))
        api_versions[v] = api_versions.get(v, 0) + 1
    if api_versions:
        av_rows = [[_esc(v), str(c)] for v, c in sorted(api_versions.items())]
        arch_sections.append(
            f'<div class="card"><h2>API Version Distribution</h2>'
            f'{_table_html(["API Version", "Count"], av_rows)}</div>'
        )
    # Automation overlap
    auto_objects = {}
    for item in data.get("trigger_findings", []):
        obj = item.get("object", "")
        if obj:
            auto_objects.setdefault(obj, set()).add("Trigger")
    for item in data.get("flow_scores", []):
        obj = item.get("object", "")
        if obj:
            auto_objects.setdefault(obj, set()).add("Flow")
    for item in data.get("process_builders", []):
        obj = item.get("object", "")
        if obj:
            auto_objects.setdefault(obj, set()).add("Process Builder")
    for item in data.get("workflow_rules", []):
        obj = item.get("object", "")
        if obj:
            auto_objects.setdefault(obj, set()).add("Workflow Rule")
    overlaps = {k: v for k, v in auto_objects.items() if len(v) > 1}
    if overlaps:
        ol_rows = [[_esc(obj), _esc(", ".join(sorted(types)))] for obj, types in sorted(overlaps.items())]
        arch_sections.append(
            f'<div class="card"><h2>Automation Overlap ({len(overlaps)} objects)</h2>'
            f'{_table_html(["Object", "Automation Types"], ol_rows)}</div>'
        )
    if arch_sections:
        body = "\n".join(arch_sections)
        p = out / "Architectural_Analysis.html"
        p.write_text(_standalone_html("Architectural Analysis", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    # ── 10. Customer Report (DOCX) ──
    has_content = bool(generated) or summary.get("overall_score", 0) > 0
    if HAS_DOCX and has_content:
        doc = docx.Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        title = doc.add_heading("Salesforce Health Assessment", level=0)
        for run in title.runs:
            run.font.color.rgb = _hex_to_rgb(BRAND_BLUE)
        doc.add_paragraph(f"Prepared for {org_name} | {run_date}")

        overall = summary["overall_score"]
        rating = summary["overall_rating"]
        p = doc.add_paragraph()
        run = p.add_run(f"Overall Health Score: {overall:.0f} / 100 — {rating}")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_heading("Key Findings", level=1)
        # Top risks in business language
        sev_counts = summary["severity_counts"]
        if sev_counts.get("CRITICAL", 0):
            doc.add_paragraph(
                f'{sev_counts["CRITICAL"]} critical issue(s) require immediate attention '
                f'to protect data integrity and security.',
                style="List Bullet",
            )
        if sev_counts.get("HIGH", 0):
            doc.add_paragraph(
                f'{sev_counts["HIGH"]} high-priority issue(s) should be addressed '
                f'within the next 30 days.',
                style="List Bullet",
            )
        if summary.get("licensing", {}).get("underutilised", 0):
            doc.add_paragraph(
                f'{summary["licensing"]["underutilised"]} licence(s) are underutilised — '
                f'review for potential cost savings.',
                style="List Bullet",
            )
        if summary.get("team", {}).get("inactive_90d", 0):
            doc.add_paragraph(
                f'{summary["team"]["inactive_90d"]} user(s) have not logged in for 90+ days.',
                style="List Bullet",
            )

        doc.add_heading("Recommendations", level=1)
        for rec in _recommendations_list(data)[:5]:
            # Strip domain prefix brackets for customer-facing
            clean = rec.split("] ", 1)[-1] if "]" in rec else rec
            doc.add_paragraph(clean, style="List Number")

        p = doc.add_paragraph()
        run = p.add_run(f"\nGenerated by Salesforce Core Skills | {run_date}")
        run.font.size = Pt(9)
        run.font.color.rgb = _hex_to_rgb(MUTED)

        p = out / "Customer_Report.docx"
        doc.save(p)
        generated.append(str(p))

    # ── 11. Strategic Engagement Plan (DOCX) ──
    if HAS_DOCX and has_content:
        doc = docx.Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        title = doc.add_heading("Strategic Engagement Plan", level=0)
        for run in title.runs:
            run.font.color.rgb = _hex_to_rgb(BRAND_BLUE)
        doc.add_paragraph(f"{org_name} | {run_date}")

        sev_counts = summary["severity_counts"]

        doc.add_heading("Phase 1 — Critical (0–30 days)", level=1)
        doc.add_paragraph(
            f'{sev_counts.get("CRITICAL", 0)} critical finding(s) to remediate immediately.'
        )
        for rec in _recommendations_list(data):
            if any(w in rec.lower() for w in ("critical", "modifyalldata", "security")):
                doc.add_paragraph(rec, style="List Bullet")

        doc.add_heading("Phase 2 — High Priority (30–90 days)", level=1)
        doc.add_paragraph(
            f'{sev_counts.get("HIGH", 0)} high-priority finding(s).'
        )

        doc.add_heading("Phase 3 — Improvements (90–180 days)", level=1)
        doc.add_paragraph(
            f'{sev_counts.get("MEDIUM", 0)} medium-priority and '
            f'{sev_counts.get("LOW", 0)} low-priority finding(s).'
        )

        doc.add_heading("Workstreams", level=1)
        for ws_name in ["Security & Governance", "Code Quality", "Automation Modernisation",
                        "Data Model & Quality", "Licensing Optimisation"]:
            doc.add_paragraph(ws_name, style="List Bullet")

        p = doc.add_paragraph()
        run = p.add_run(f"\nGenerated by Salesforce Core Skills | {run_date}")
        run.font.size = Pt(9)
        run.font.color.rgb = _hex_to_rgb(MUTED)

        p = out / "Strategic_Engagement_Plan.docx"
        doc.save(p)
        generated.append(str(p))

    # ── 12. Technical Team Briefing (HTML) ──
    briefing_sections = []
    # Worst-scoring components
    worst = []
    for domain, key, max_s in [("Apex", "apex_scores", 150), ("Flows", "flow_scores", 110),
                                ("LWC", "lwc_scores", 165), ("Metadata", "metadata_scores", 120)]:
        for item in data.get(key, []):
            s = item.get("score", 0)
            m = item.get("max_score", max_s)
            pct = (s / m * 100) if m > 0 else 0
            if pct < 70:
                worst.append((pct, item.get("name", ""), domain,
                              "; ".join(x if isinstance(x, str) else x.get("message", "")
                                        for x in item.get("issues", [])[:2])))
    worst.sort(key=lambda x: x[0])
    if worst:
        rows = [[_esc(n), _esc(d), f"{p:.0f}%", _esc(iss)] for p, n, d, iss in worst[:15]]
        briefing_sections.append(
            f'<div class="card"><h2>Code Quality Hotspots</h2>'
            f'{_table_html(["Component", "Domain", "Score", "Top Issues"], rows)}</div>'
        )
    # Coverage gaps
    gaps = [c for c in tc.get("classes", []) if c.get("coverage_pct", 0) < 75]
    if gaps:
        rows = [[_esc(c.get("name", "")), f'{c.get("coverage_pct", 0):.1f}%']
                for c in sorted(gaps, key=lambda x: x.get("coverage_pct", 0))]
        briefing_sections.append(
            f'<div class="card"><h2>Test Coverage Gaps ({len(gaps)} classes below 75%)</h2>'
            f'{_table_html(["Class", "Coverage"], rows)}</div>'
        )
    if briefing_sections:
        body = "\n".join(briefing_sections)
        p = out / "Technical_Team_Briefing.html"
        p.write_text(_standalone_html("Technical Team Briefing", org_name, run_date, body), encoding="utf-8")
        generated.append(str(p))

    return generated


# ── Main ────────────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(description="Generate Salesforce Org Audit Reports")
    parser.add_argument("--input-dir", required=True, help="Directory containing scored JSON files")
    parser.add_argument("--output-dir", required=True, help="Directory for generated reports")
    parser.add_argument("--org-name", required=True, help="Salesforce org name")
    parser.add_argument("--org-id", default="", help="18-char Org ID")
    parser.add_argument("--instance", default="", help="Salesforce instance (e.g. CS42)")
    parser.add_argument("--run-date", default=None, help="Audit date (YYYY-MM-DD). Defaults to today.")
    parser.add_argument(
        "--standalone",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Generate standalone reports",
    )
    args = parser.parse_args()

    run_date = args.run_date or date.today().isoformat()

    # Validate input dir
    input_path = Path(args.input_dir)
    if not input_path.is_dir():
        print(f"ERROR: Input directory does not exist: {args.input_dir}", file=sys.stderr)
        sys.exit(1)

    # Create output dir
    output_path = Path(args.output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    # Load and compute
    data = load_inputs(args.input_dir)
    summary = compute_summary(data)

    # Override org info from CLI args (takes precedence over counts.json)
    org_name = args.org_name
    org_id = args.org_id or summary.get("org_id", "")
    instance = args.instance or summary.get("instance", "")

    generated = []

    # HTML
    html_path = output_path / "Salesforce_Org_Audit_Report.html"
    generate_html(data, summary, org_name, org_id, instance, run_date, html_path)
    generated.append(str(html_path))
    print(f"  HTML:  {html_path}")

    # DOCX
    docx_path = output_path / "Salesforce_Org_Audit_Report.docx"
    result = generate_docx(data, summary, org_name, org_id, instance, run_date, docx_path)
    if result:
        generated.append(str(docx_path))
        print(f"  DOCX:  {docx_path}")

    # XLSX
    xlsx_path = output_path / "Salesforce_Org_Audit_Scores.xlsx"
    result = generate_xlsx(data, summary, org_name, org_id, instance, run_date, xlsx_path)
    if result:
        generated.append(str(xlsx_path))
        print(f"  XLSX:  {xlsx_path}")

    # JSON summary
    json_path = output_path / "audit_summary.json"
    generate_json_summary(summary, run_date, json_path)
    generated.append(str(json_path))
    print(f"  JSON:  {json_path}")

    # Standalone reports
    if args.standalone:
        standalone = generate_standalone_reports(
            data, summary, org_name, run_date, output_path,
        )
        generated.extend(standalone)
        for p in standalone:
            print(f"  STANDALONE:  {p}")

    print(f"\nOverall Score: {summary['overall_score']:.0f}/100 — {summary['overall_rating']}")
    print(f"Generated {len(generated)} report(s)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
